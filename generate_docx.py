#!/usr/bin/env python3
"""
Generate .docx files for 12 Stripe Sessions research companies.
Reads .md files, parses DATABASE FIELDS, produces formatted .docx.
"""

import re
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE = "/Users/germantatis/Desktop/GTMCoding/Stripe Sessions 2026/Research"

COMPANIES = [
    ("1Password", "1password-stripe-sessions-research"),
    ("Ancestry", "ancestry-stripe-sessions-research"),
    ("Calendly", "calendly-stripe-sessions-research"),
    ("Hostinger", "hostinger-stripe-sessions-research"),
    ("LegalZoom", "legalzoom-stripe-sessions-research"),
    ("Lululemon Athletica", "lululemon-stripe-sessions-research"),
    ("MongoDB", "mongodb-stripe-sessions-research"),
    ("Rippling", "rippling-stripe-sessions-research"),
    ("SeatGeek", "seatgeek-stripe-sessions-research"),
    ("Ssense", "ssense-stripe-sessions-research"),
    ("Zillow", "zillow-stripe-sessions-research"),
    ("hims & hers", "hims-hers-stripe-sessions-research"),
]

PURPLE = RGBColor(0x5B, 0x4C, 0xFF)
DARK = RGBColor(0x1A, 0x1A, 0x2E)
GRAY_TEXT = RGBColor(0x99, 0x99, 0x99)
LIGHT_GRAY = RGBColor(0xCC, 0xCC, 0xCC)
FIELD_BG = "F2F2F2"


def set_cell_shading(cell, color_hex):
    """Set background shading on a table cell."""
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color_hex)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def set_paragraph_shading(paragraph, color_hex):
    """Set background shading on a paragraph."""
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color_hex)
    shading.set(qn("w:val"), "clear")
    pPr = paragraph._element.get_or_add_pPr()
    pPr.append(shading)


def add_run(paragraph, text, font_name="Calibri", size=11, bold=False, color=None, italic=False):
    """Add a formatted run to a paragraph."""
    run = paragraph.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color
    return run


def add_separator(doc):
    """Add a light gray separator line."""
    p = doc.add_paragraph()
    add_run(p, "─" * 58, color=LIGHT_GRAY, size=10)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)


def add_double_separator(doc):
    """Add a double-line separator."""
    p = doc.add_paragraph()
    add_run(p, "═" * 39, size=10, bold=True)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)


def parse_md(md_path):
    """Parse .md file and extract all fields."""
    with open(md_path, "r", encoding="utf-8") as f:
        content = f.read()

    # Extract the code block with DATABASE FIELDS
    code_match = re.search(r"```(.*?)```", content, re.DOTALL)
    if not code_match:
        raise ValueError(f"No code block found in {md_path}")
    db_block = code_match.group(1)

    # Extract the research backing section (everything after the code block)
    research_section = content[code_match.end():]

    data = {}

    # Company name from DATABASE FIELDS line
    m = re.search(r"DATABASE FIELDS:\s*(.+)", db_block)
    data["db_name"] = m.group(1).strip() if m else ""

    # Logo
    m = re.search(r"Logo:\s*(.+)", db_block)
    data["logo"] = m.group(1).strip() if m else ""

    # Nombre merchant
    m = re.search(r"Nombre merchant:\s*(.+)", db_block)
    data["nombre"] = m.group(1).strip() if m else ""

    # Pain Point titles
    data["pp_titles"] = []
    for i in range(1, 6):
        m = re.search(rf"Tittle_Pain Point_{i}:\s*(.+)", db_block)
        data["pp_titles"].append(m.group(1).strip() if m else "")

    # Pain Point descriptions
    data["pp_descs"] = []
    for i in range(1, 6):
        m = re.search(rf"Desc_Pain Point_{i}:\s*(.+)", db_block)
        data["pp_descs"].append(m.group(1).strip() if m else "")

    # PSPs
    data["psps"] = []
    for i in range(1, 10):
        m = re.search(rf"PSP_{i}:\s*(.+)", db_block)
        if m:
            data["psps"].append(m.group(1).strip())
        else:
            break

    # Local methods
    data["local_methods"] = []
    for i in range(1, 15):
        m = re.search(rf"Local_M_{i}:\s*(.+)", db_block)
        if m:
            data["local_methods"].append(m.group(1).strip())
        else:
            break

    # Yuno capability titles
    data["yuno_titles"] = []
    for i in range(1, 5):
        m = re.search(rf"Tittle_Yuno_Cap{i}:\s*(.+)", db_block)
        data["yuno_titles"].append(m.group(1).strip() if m else "")

    # Yuno capability descriptions
    data["yuno_descs"] = []
    for i in range(1, 5):
        m = re.search(rf"Desc_Yuno_Cap{i}:\s*(.+)", db_block)
        data["yuno_descs"].append(m.group(1).strip() if m else "")

    # Research backing content
    data["research"] = research_section.strip()

    return data


def build_docx(data, output_path):
    """Build the formatted .docx from parsed data."""
    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)

    # ========== TITLE ==========
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_run(title_p, f"{data['db_name'].upper()} | Stripe Sessions Research (SS Database Format)",
            size=16, bold=True, color=PURPLE)
    title_p.paragraph_format.space_after = Pt(4)

    # Date
    date_p = doc.add_paragraph()
    add_run(date_p, "Generated: 2026-04-21", size=10, italic=True, color=GRAY_TEXT)
    date_p.paragraph_format.space_after = Pt(12)

    # ========== DATABASE FIELDS HEADER ==========
    add_double_separator(doc)

    db_header = doc.add_paragraph()
    add_run(db_header, f"DATABASE FIELDS: {data['db_name']}", size=14, bold=True, color=PURPLE)
    db_header.paragraph_format.space_before = Pt(2)
    db_header.paragraph_format.space_after = Pt(2)

    add_double_separator(doc)
    doc.add_paragraph()

    # ========== LOGO & NOMBRE ==========
    logo_p = doc.add_paragraph()
    set_paragraph_shading(logo_p, FIELD_BG)
    add_run(logo_p, "Logo: ", size=10, bold=True, color=PURPLE)
    add_run(logo_p, data["logo"], size=10)
    add_run(logo_p, f"  [{len(data['logo'])} chars]", size=8, color=GRAY_TEXT, italic=True)

    nombre_p = doc.add_paragraph()
    set_paragraph_shading(nombre_p, FIELD_BG)
    add_run(nombre_p, "Nombre merchant: ", size=10, bold=True, color=PURPLE)
    add_run(nombre_p, data["nombre"], size=10)
    add_run(nombre_p, f"  [{len(data['nombre'])} chars]", size=8, color=GRAY_TEXT, italic=True)

    doc.add_paragraph()

    # ========== SLIDE 1: PAIN POINTS ==========
    section_p = doc.add_paragraph()
    add_run(section_p, "SLIDE 1: PAIN POINTS", size=13, bold=True, color=PURPLE)
    section_p.paragraph_format.space_after = Pt(6)

    # Pain Point Titles
    for i, title in enumerate(data["pp_titles"], 1):
        p = doc.add_paragraph()
        set_paragraph_shading(p, FIELD_BG)
        add_run(p, f"Tittle_Pain Point_{i}: ", size=10, bold=True, color=PURPLE)
        add_run(p, title, size=10)
        add_run(p, f"  [{len(title)} chars]", size=8, color=GRAY_TEXT, italic=True)

    doc.add_paragraph()

    # Pain Point Descriptions
    for i, desc in enumerate(data["pp_descs"], 1):
        p = doc.add_paragraph()
        set_paragraph_shading(p, FIELD_BG)
        add_run(p, f"Desc_Pain Point_{i}: ", size=10, bold=True, color=PURPLE)
        add_run(p, desc, size=10)
        p2 = doc.add_paragraph()
        set_paragraph_shading(p2, FIELD_BG)
        add_run(p2, f"[{len(desc)} chars]", size=8, color=GRAY_TEXT, italic=True)
        p2.paragraph_format.space_before = Pt(0)
        p2.paragraph_format.space_after = Pt(6)

    doc.add_paragraph()

    # ========== SLIDE 1: PSP TOPOLOGY ==========
    section_p2 = doc.add_paragraph()
    add_run(section_p2, "SLIDE 1: PSP TOPOLOGY", size=13, bold=True, color=PURPLE)
    section_p2.paragraph_format.space_after = Pt(6)

    # PSP Table
    if data["psps"]:
        table = doc.add_table(rows=len(data["psps"]) + 1, cols=2)
        table.style = "Table Grid"

        # Header row
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = ""
        hdr_cells[1].text = ""
        hdr_p0 = hdr_cells[0].paragraphs[0]
        hdr_p0.clear()
        add_run(hdr_p0, "Field", size=10, bold=True, color=PURPLE)
        set_cell_shading(hdr_cells[0], "E8E5FF")

        hdr_p1 = hdr_cells[1].paragraphs[0]
        hdr_p1.clear()
        add_run(hdr_p1, "Value", size=10, bold=True, color=PURPLE)
        set_cell_shading(hdr_cells[1], "E8E5FF")

        for idx, psp in enumerate(data["psps"]):
            row = table.rows[idx + 1]
            c0 = row.cells[0]
            c1 = row.cells[1]
            c0_p = c0.paragraphs[0]
            c0_p.clear()
            add_run(c0_p, f"PSP_{idx+1}", size=10, bold=True)
            c1_p = c1.paragraphs[0]
            c1_p.clear()
            add_run(c1_p, psp, size=10)
            if idx % 2 == 1:
                set_cell_shading(c0, FIELD_BG)
                set_cell_shading(c1, FIELD_BG)

    doc.add_paragraph()

    # ========== SLIDE 1: MISSING LOCAL METHODS ==========
    section_p3 = doc.add_paragraph()
    add_run(section_p3, "SLIDE 1: MISSING LOCAL METHODS (BLIND SPOTS)", size=13, bold=True, color=PURPLE)
    section_p3.paragraph_format.space_after = Pt(6)

    # Local Methods Table
    if data["local_methods"]:
        table2 = doc.add_table(rows=len(data["local_methods"]) + 1, cols=2)
        table2.style = "Table Grid"

        hdr2_cells = table2.rows[0].cells
        hdr2_p0 = hdr2_cells[0].paragraphs[0]
        hdr2_p0.clear()
        add_run(hdr2_p0, "Field", size=10, bold=True, color=PURPLE)
        set_cell_shading(hdr2_cells[0], "E8E5FF")

        hdr2_p1 = hdr2_cells[1].paragraphs[0]
        hdr2_p1.clear()
        add_run(hdr2_p1, "Value", size=10, bold=True, color=PURPLE)
        set_cell_shading(hdr2_cells[1], "E8E5FF")

        for idx, method in enumerate(data["local_methods"]):
            row = table2.rows[idx + 1]
            c0 = row.cells[0]
            c1 = row.cells[1]
            c0_p = c0.paragraphs[0]
            c0_p.clear()
            add_run(c0_p, f"Local_M_{idx+1}", size=10, bold=True)
            c1_p = c1.paragraphs[0]
            c1_p.clear()
            add_run(c1_p, method, size=10)
            if idx % 2 == 1:
                set_cell_shading(c0, FIELD_BG)
                set_cell_shading(c1, FIELD_BG)

    doc.add_paragraph()

    # ========== SLIDE 2: YUNO CAPABILITIES ==========
    section_p4 = doc.add_paragraph()
    add_run(section_p4, "SLIDE 2: YUNO CAPABILITIES", size=13, bold=True, color=PURPLE)
    section_p4.paragraph_format.space_after = Pt(6)

    # Yuno Capability Titles
    for i, title in enumerate(data["yuno_titles"], 1):
        p = doc.add_paragraph()
        set_paragraph_shading(p, FIELD_BG)
        add_run(p, f"Tittle_Yuno_Cap{i}: ", size=10, bold=True, color=PURPLE)
        add_run(p, title, size=10)
        add_run(p, f"  [{len(title)} chars]", size=8, color=GRAY_TEXT, italic=True)

    doc.add_paragraph()

    # Yuno Capability Descriptions
    for i, desc in enumerate(data["yuno_descs"], 1):
        p = doc.add_paragraph()
        set_paragraph_shading(p, FIELD_BG)
        add_run(p, f"Desc_Yuno_Cap{i}: ", size=10, bold=True, color=PURPLE)
        add_run(p, desc, size=10)
        p2 = doc.add_paragraph()
        set_paragraph_shading(p2, FIELD_BG)
        add_run(p2, f"[{len(desc)} chars]", size=8, color=GRAY_TEXT, italic=True)
        p2.paragraph_format.space_before = Pt(0)
        p2.paragraph_format.space_after = Pt(6)

    # ========== PAGE BREAK ==========
    doc.add_page_break()

    # ========== RESEARCH BACKING ==========
    research_header = doc.add_paragraph()
    add_run(research_header, "RESEARCH BACKING | Key Facts for the Meeting", size=14, bold=True, color=PURPLE)
    research_header.paragraph_format.space_after = Pt(8)

    add_separator(doc)
    doc.add_paragraph()

    # Parse and render research section
    render_research(doc, data["research"])

    doc.save(output_path)
    print(f"  Saved: {output_path}")


def render_research(doc, research_text):
    """Render the research backing section from markdown to docx."""
    lines = research_text.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Skip markdown separators and empty heading markers
        if stripped in ("---", ""):
            i += 1
            continue

        # Heading 2 (## ...)
        if stripped.startswith("## "):
            # Skip the main "RESEARCH BACKING" heading since we already added it
            if "RESEARCH BACKING" in stripped:
                i += 1
                continue
            p = doc.add_paragraph()
            text = stripped.lstrip("# ").strip()
            add_run(p, text, size=13, bold=True, color=PURPLE)
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(4)
            i += 1
            continue

        # Bold section header (**text:**)
        bold_match = re.match(r"^\*\*(.+?)\*\*\s*$", stripped)
        if bold_match:
            p = doc.add_paragraph()
            add_run(p, bold_match.group(1), size=11, bold=True)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(2)
            i += 1
            continue

        # Bullet point
        if stripped.startswith("- "):
            bullet_text = stripped[2:].strip()
            # Remove markdown links: [text](url) -> text
            bullet_text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", bullet_text)
            # Remove bold markers
            bullet_text = re.sub(r"\*\*(.+?)\*\*", r"\1", bullet_text)
            p = doc.add_paragraph(style="List Bullet")
            add_run(p, bullet_text, size=10)
            i += 1
            continue

        # Numbered items (1. **text** | desc)
        numbered_match = re.match(r"^\d+\.\s+(.+)$", stripped)
        if numbered_match:
            text = numbered_match.group(1)
            text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
            p = doc.add_paragraph()
            add_run(p, text, size=10)
            p.paragraph_format.space_before = Pt(2)
            i += 1
            continue

        # Market analysis blocks
        market_match = re.match(r"^MARKET \d+:", stripped)
        if market_match:
            p = doc.add_paragraph()
            add_run(p, stripped, size=10, bold=True)
            p.paragraph_format.space_before = Pt(4)
            i += 1
            continue

        # Indented lines (like Accepted:, Missing:, Note:)
        if stripped.startswith("Accepted:") or stripped.startswith("Missing:") or stripped.startswith("Note:") or stripped.startswith("Insight:") or stripped.startswith("Renters"):
            p = doc.add_paragraph()
            add_run(p, "  " + stripped, size=10)
            i += 1
            continue

        # Table rows (markdown tables)
        if stripped.startswith("|"):
            # Collect all table rows
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i].strip())
                i += 1

            # Parse table
            if len(table_lines) >= 2:
                # Header
                header_cells = [c.strip() for c in table_lines[0].split("|") if c.strip()]
                # Skip separator row
                data_rows = []
                for tl in table_lines[2:]:
                    cells = [c.strip() for c in tl.split("|") if c.strip()]
                    if cells:
                        data_rows.append(cells)

                if header_cells and data_rows:
                    num_cols = len(header_cells)
                    table = doc.add_table(rows=len(data_rows) + 1, cols=num_cols)
                    table.style = "Table Grid"

                    for ci, hc in enumerate(header_cells):
                        hp = table.rows[0].cells[ci].paragraphs[0]
                        hp.clear()
                        add_run(hp, hc, size=9, bold=True, color=PURPLE)
                        set_cell_shading(table.rows[0].cells[ci], "E8E5FF")

                    for ri, dr in enumerate(data_rows):
                        for ci in range(min(num_cols, len(dr))):
                            cp = table.rows[ri + 1].cells[ci].paragraphs[0]
                            cp.clear()
                            add_run(cp, dr[ci], size=9)
                            if ri % 2 == 1:
                                set_cell_shading(table.rows[ri + 1].cells[ci], FIELD_BG)

                doc.add_paragraph()
            continue

        # Regular text
        if stripped:
            # Remove markdown formatting
            clean = re.sub(r"\*\*(.+?)\*\*", r"\1", stripped)
            clean = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", clean)
            p = doc.add_paragraph()
            add_run(p, clean, size=10)

        i += 1


def main():
    success = 0
    errors = []

    for folder_name, file_stem in COMPANIES:
        md_path = os.path.join(BASE, folder_name, f"{file_stem}.md")
        docx_path = os.path.join(BASE, folder_name, f"{file_stem}.docx")

        print(f"\nProcessing: {folder_name}")

        if not os.path.exists(md_path):
            print(f"  ERROR: .md file not found: {md_path}")
            errors.append(folder_name)
            continue

        try:
            data = parse_md(md_path)
            build_docx(data, docx_path)
            success += 1
        except Exception as e:
            print(f"  ERROR: {e}")
            errors.append(folder_name)

    print(f"\n{'='*50}")
    print(f"Completed: {success}/12 files generated successfully")
    if errors:
        print(f"Errors: {', '.join(errors)}")


if __name__ == "__main__":
    main()
