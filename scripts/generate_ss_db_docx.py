#!/usr/bin/env python3
"""
Generates .docx files from Stripe Sessions Research .md files (SS Database Format).
Loops through 13 companies that have .md files but no .docx yet.
"""

import os
import re
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

RESEARCH_DIR = "/Users/germantatis/Desktop/GTMCoding/Stripe Sessions 2026/Research"

PURPLE = RGBColor(0x5B, 0x4C, 0xFF)
GRAY = RGBColor(0x99, 0x99, 0x99)
DARK = RGBColor(0x1A, 0x1A, 0x2E)

# 13 companies: (folder_name, file_slug)
COMPANIES = [
    ("Alibaba Group", "alibaba-group-stripe-sessions-research"),
    ("DoorDash", "doordash-stripe-sessions-research"),
    ("Electronic Arts", "electronic-arts-stripe-sessions-research"),
    ("Epic Games", "epic-games-stripe-sessions-research"),
    ("Google", "google-stripe-sessions-research"),
    ("Marriott International", "marriott-international-stripe-sessions-research"),
    ("Microsoft", "microsoft-stripe-sessions-research"),
    ("Netflix", "netflix-stripe-sessions-research"),
    ("Nike", "nike-stripe-sessions-research"),
    ("Roblox", "roblox-stripe-sessions-research"),
    ("Sephora", "sephora-stripe-sessions-research"),
    ("Spotify", "spotify-stripe-sessions-research"),
    ("Walmart", "walmart-stripe-sessions-research"),
]


def parse_database_fields(md_text):
    """Extract all DATABASE FIELDS from the code block in the .md file."""
    data = {}

    # Extract company name from the title line
    m = re.search(r'^# (.+?) \|', md_text, re.MULTILINE)
    data['company'] = m.group(1).strip() if m else 'COMPANY'

    # Extract date
    m = re.search(r'\*Generated: (.+?)\*', md_text)
    data['date'] = m.group(1).strip() if m else ''

    # Extract the code block containing DATABASE FIELDS
    code_match = re.search(r'```\n(.*?)```', md_text, re.DOTALL)
    if not code_match:
        return data
    code_block = code_match.group(1)

    # Logo
    m = re.search(r'Logo:\s*(.+)', code_block)
    data['logo'] = m.group(1).strip() if m else ''

    # Nombre merchant
    m = re.search(r'Nombre merchant:\s*(.+)', code_block)
    data['nombre'] = m.group(1).strip() if m else ''

    # Pain Point titles and descriptions
    data['pain_titles'] = {}
    data['pain_descs'] = {}
    for i in range(1, 6):
        m = re.search(rf'Tittle_Pain Point_{i}:\s*(.+)', code_block)
        data['pain_titles'][i] = m.group(1).strip() if m else ''
        m = re.search(rf'Desc_Pain Point_{i}:\s*(.+)', code_block)
        data['pain_descs'][i] = m.group(1).strip() if m else ''

    # PSPs
    data['psps'] = {}
    for i in range(1, 5):
        m = re.search(rf'PSP_{i}:\s*(.+)', code_block)
        data['psps'][i] = m.group(1).strip() if m else ''

    # Local methods
    data['local_methods'] = {}
    for i in range(1, 11):
        m = re.search(rf'Local_M_{i}:\s*(.+)', code_block)
        data['local_methods'][i] = m.group(1).strip() if m else ''

    # Yuno capabilities titles and descriptions
    data['yuno_titles'] = {}
    data['yuno_descs'] = {}
    for i in range(1, 5):
        m = re.search(rf'Tittle_Yuno_Cap{i}:\s*(.+)', code_block)
        data['yuno_titles'][i] = m.group(1).strip() if m else ''
        m = re.search(rf'Desc_Yuno_Cap{i}:\s*(.+)', code_block)
        data['yuno_descs'][i] = m.group(1).strip() if m else ''

    # Research backing: everything after the code block closing
    after_code = md_text.split('```', 2)
    if len(after_code) >= 3:
        data['research_text'] = after_code[2].strip()
    else:
        data['research_text'] = ''

    return data


def add_heading2_purple(doc, text):
    """Add a Heading 2 with purple color."""
    p = doc.add_heading(text, level=2)
    for run in p.runs:
        run.font.color.rgb = PURPLE
        run.font.name = 'Calibri'
    return p


def add_field_pair(doc, label, value, char_limit=None):
    """Add a field label (bold) and value. Optionally show char count."""
    p = doc.add_paragraph()
    r1 = p.add_run(f"{label}: ")
    r1.font.bold = True
    r1.font.name = 'Calibri'
    r1.font.size = Pt(10)

    display_value = value
    if char_limit:
        count = len(value)
        display_value = f"{value}  [{count}/{char_limit} chars]"

    r2 = p.add_run(display_value)
    r2.font.name = 'Calibri'
    r2.font.size = Pt(10)
    return p


def add_desc_field(doc, label, value):
    """Add a description field with char count indicator and gray background."""
    p = doc.add_paragraph()
    r1 = p.add_run(f"{label}: ")
    r1.font.bold = True
    r1.font.name = 'Calibri'
    r1.font.size = Pt(10)

    count = len(value)
    r2 = p.add_run(f"{value}")
    r2.font.name = 'Calibri'
    r2.font.size = Pt(10)

    # Add char count indicator
    r3 = p.add_run(f"  [{count} chars]")
    r3.font.name = 'Calibri'
    r3.font.size = Pt(8)
    r3.font.color.rgb = GRAY
    r3.font.italic = True
    return p


def add_simple_field(doc, label, value):
    """Add a simple field line."""
    p = doc.add_paragraph()
    r1 = p.add_run(f"{label}: ")
    r1.font.bold = True
    r1.font.name = 'Calibri'
    r1.font.size = Pt(10)

    r2 = p.add_run(value)
    r2.font.name = 'Calibri'
    r2.font.size = Pt(10)
    return p


def build_psp_table(doc, psps):
    """Build a simple table showing the PSP topology."""
    table = doc.add_table(rows=1, cols=len(psps))
    table.style = 'Light Grid Accent 1'

    # Header row
    for i, (idx, psp) in enumerate(psps.items()):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        r = p.add_run(f"PSP {idx}: {psp}")
        r.font.name = 'Calibri'
        r.font.size = Pt(10)
        r.font.bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return table


def build_gap_table(doc, local_methods):
    """Build a table showing the blind spot / local methods."""
    # 2 rows of 5
    active = {k: v for k, v in local_methods.items() if v}
    items = list(active.values())

    rows_needed = (len(items) + 4) // 5  # ceil division
    cols = min(len(items), 5)
    if cols == 0:
        return

    table = doc.add_table(rows=rows_needed, cols=cols)
    table.style = 'Light Grid Accent 1'

    idx = 0
    for row_i in range(rows_needed):
        for col_i in range(cols):
            if idx < len(items):
                cell = table.rows[row_i].cells[col_i]
                cell.text = ''
                p = cell.paragraphs[0]
                r = p.add_run(items[idx])
                r.font.name = 'Calibri'
                r.font.size = Pt(9)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                idx += 1
    return table


def parse_research_backing(research_text):
    """Parse the research backing markdown into structured sections."""
    sections = []
    current_section = None
    current_items = []

    for line in research_text.split('\n'):
        stripped = line.strip()
        if not stripped:
            continue

        # Section headers: ## heading or **heading:**
        if stripped.startswith('## '):
            if current_section:
                sections.append((current_section, current_items))
            current_section = stripped.lstrip('# ').strip()
            current_items = []
        elif stripped.startswith('**') and stripped.endswith('**'):
            if current_section:
                sections.append((current_section, current_items))
            current_section = stripped.replace('**', '').strip()
            current_items = []
        elif stripped.startswith('- ') or stripped.startswith('* '):
            # Clean markdown formatting from bullet items
            item = stripped.lstrip('- *').strip()
            # Remove bold markers
            item = re.sub(r'\*\*(.+?)\*\*', r'\1', item)
            # Remove link formatting: [text](url) -> text (url)
            item = re.sub(r'\[(.+?)\]\((.+?)\)', r'\1 (\2)', item)
            current_items.append(item)
        elif re.match(r'^\d+\.\s', stripped):
            item = re.sub(r'^\d+\.\s*', '', stripped)
            item = re.sub(r'\*\*(.+?)\*\*', r'\1', item)
            item = re.sub(r'\[(.+?)\]\((.+?)\)', r'\1 (\2)', item)
            current_items.append(item)
        elif stripped.startswith('MARKET '):
            # Market gap analysis lines
            current_items.append(stripped)
        else:
            # Regular text
            cleaned = re.sub(r'\*\*(.+?)\*\*', r'\1', stripped)
            cleaned = re.sub(r'\[(.+?)\]\((.+?)\)', r'\1 (\2)', cleaned)
            if cleaned and cleaned != '---':
                current_items.append(cleaned)

    if current_section:
        sections.append((current_section, current_items))

    return sections


def build_docx(data, docx_path):
    """Build the .docx file from parsed data."""
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10)

    # ===== TITLE =====
    p = doc.add_heading(f"{data['company'].upper()} | Stripe Sessions Research (SS Database Format)", level=1)
    for run in p.runs:
        run.font.color.rgb = DARK
        run.font.name = 'Calibri'

    p = doc.add_paragraph()
    r = p.add_run(f"Generated: {data['date']}")
    r.font.name = 'Calibri'
    r.font.size = Pt(10)
    r.font.italic = True
    r.font.color.rgb = GRAY

    # ===== DATABASE FIELDS HEADER =====
    p = doc.add_paragraph()
    r = p.add_run(f"DATABASE FIELDS: {data.get('nombre', data['company'])}")
    r.font.name = 'Calibri'
    r.font.size = Pt(14)
    r.font.bold = True
    r.font.color.rgb = PURPLE

    doc.add_paragraph()

    # Logo and Nombre
    add_simple_field(doc, "Logo", data.get('logo', ''))
    add_simple_field(doc, "Nombre merchant", data.get('nombre', ''))

    # ===== SLIDE 1: PAIN POINTS =====
    add_heading2_purple(doc, "SLIDE 1: PAIN POINTS")

    for i in range(1, 6):
        title = data['pain_titles'].get(i, '')
        desc = data['pain_descs'].get(i, '')
        if title:
            add_field_pair(doc, f"Tittle_Pain Point_{i}", title)
        if desc:
            add_desc_field(doc, f"Desc_Pain Point_{i}", desc)
        if title or desc:
            doc.add_paragraph()  # blank line between pairs

    # ===== SLIDE 1: PSP TOPOLOGY =====
    add_heading2_purple(doc, "SLIDE 1: PSP TOPOLOGY")

    # Show as individual fields
    for i in range(1, 5):
        psp = data['psps'].get(i, '')
        if psp:
            add_simple_field(doc, f"PSP_{i}", psp)

    doc.add_paragraph()

    # PSP topology table
    active_psps = {k: v for k, v in data['psps'].items() if v}
    if active_psps:
        build_psp_table(doc, active_psps)

    # ===== SLIDE 1: MISSING LOCAL METHODS =====
    add_heading2_purple(doc, "SLIDE 1: MISSING LOCAL METHODS (BLIND SPOTS)")

    for i in range(1, 11):
        method = data['local_methods'].get(i, '')
        if method:
            add_simple_field(doc, f"Local_M_{i}", method)

    doc.add_paragraph()

    # Gap analysis table
    build_gap_table(doc, data['local_methods'])

    doc.add_page_break()

    # ===== SLIDE 2: YUNO CAPABILITIES =====
    add_heading2_purple(doc, "SLIDE 2: YUNO CAPABILITIES")

    for i in range(1, 5):
        title = data['yuno_titles'].get(i, '')
        desc = data['yuno_descs'].get(i, '')
        if title:
            add_field_pair(doc, f"Tittle_Yuno_Cap{i}", title)
        if desc:
            add_desc_field(doc, f"Desc_Yuno_Cap{i}", desc)
        if title or desc:
            doc.add_paragraph()

    doc.add_page_break()

    # ===== RESEARCH BACKING =====
    p = doc.add_heading("RESEARCH BACKING | Key Facts for the Meeting", level=1)
    for run in p.runs:
        run.font.color.rgb = PURPLE
        run.font.name = 'Calibri'

    research_text = data.get('research_text', '')
    if research_text:
        sections = parse_research_backing(research_text)
        for section_title, items in sections:
            if section_title.lower() == 'sources' or section_title.lower().startswith('source'):
                # Sources section
                p = doc.add_heading(section_title, level=2)
                for run in p.runs:
                    run.font.color.rgb = PURPLE
                    run.font.name = 'Calibri'
                for item in items:
                    bp = doc.add_paragraph(item, style='List Bullet')
                    for run in bp.runs:
                        run.font.name = 'Calibri'
                        run.font.size = Pt(9)
            else:
                p = doc.add_heading(section_title, level=2)
                for run in p.runs:
                    run.font.color.rgb = PURPLE
                    run.font.name = 'Calibri'
                for item in items:
                    if item.startswith('MARKET '):
                        # Market analysis headers
                        mp = doc.add_paragraph()
                        mr = mp.add_run(item)
                        mr.font.bold = True
                        mr.font.name = 'Calibri'
                        mr.font.size = Pt(10)
                    else:
                        bp = doc.add_paragraph(item, style='List Bullet')
                        for run in bp.runs:
                            run.font.name = 'Calibri'
                            run.font.size = Pt(10)

    doc.save(docx_path)
    return docx_path


def main():
    success = []
    errors = []

    for folder_name, file_slug in COMPANIES:
        md_path = os.path.join(RESEARCH_DIR, folder_name, f"{file_slug}.md")
        docx_path = os.path.join(RESEARCH_DIR, folder_name, f"{file_slug}.docx")

        if not os.path.exists(md_path):
            errors.append(f"NOT FOUND: {md_path}")
            continue

        try:
            with open(md_path, 'r', encoding='utf-8') as f:
                md_text = f.read()

            data = parse_database_fields(md_text)
            output = build_docx(data, docx_path)
            success.append(f"OK: {folder_name} -> {os.path.basename(output)}")
            print(f"  Generated: {output}")
        except Exception as e:
            errors.append(f"ERROR: {folder_name} -> {e}")
            print(f"  FAILED: {folder_name} -> {e}")

    print(f"\n{'='*60}")
    print(f"Results: {len(success)} generated, {len(errors)} errors")
    print(f"{'='*60}")
    for s in success:
        print(f"  {s}")
    for e in errors:
        print(f"  {e}")


if __name__ == '__main__':
    main()
