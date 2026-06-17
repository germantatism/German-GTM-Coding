#!/usr/bin/env python3
"""
Converts a Stripe Sessions research .md file to a formatted .docx
Usage: python3 generate_stripe_sessions_docx.py <company-slug>
  Reads:  .../Research/<company-slug>-stripe-sessions-research.md
  Writes: .../Research/<company-slug>-stripe-sessions-research.docx
"""

import sys, os, re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.oxml.ns import qn
from lxml import etree

RESEARCH_DIR = "/Users/germantatis/Desktop/GTMCoding/Stripe Sessions 2026/Research"
PURPLE = RGBColor(0x5B, 0x4C, 0xFF)
GRAY = RGBColor(0x99, 0x99, 0x99)
DARK = RGBColor(0x1A, 0x1A, 0x2E)
ORANGE = RGBColor(0xFF, 0x66, 0x00)

def shade_paragraph(p, color='F5F5F5'):
    shd = etree.SubElement(p.paragraph_format.element, qn('w:shd'))
    shd.set(qn('w:fill'), color)
    shd.set(qn('w:val'), 'clear')

def add_slide_header(doc, text):
    p = doc.add_heading(text, level=1)
    for run in p.runs:
        run.font.size = Pt(18)
        run.font.color.rgb = PURPLE

def add_box_header(doc, label, title):
    p = doc.add_paragraph()
    r1 = p.add_run(f"{label}: ")
    r1.font.size, r1.font.bold, r1.font.color.rgb = Pt(12), True, PURPLE
    r2 = p.add_run(title)
    r2.font.size, r2.font.bold = Pt(12), True

def add_copy_block(doc, label, text):
    p = doc.add_paragraph()
    r = p.add_run(f"{label}:")
    r.font.bold, r.font.size, r.font.color.rgb = True, Pt(10), RGBColor(0x33,0x33,0x33)
    p2 = doc.add_paragraph()
    r2 = p2.add_run(text)
    r2.font.size = Pt(11)
    p2.paragraph_format.left_indent = Inches(0.3)
    shade_paragraph(p2)

def add_source(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size, r.font.italic, r.font.color.rgb = Pt(8), True, GRAY

def add_separator(doc):
    p = doc.add_paragraph()
    p.add_run('─' * 60).font.color.rgb = RGBColor(0xCC,0xCC,0xCC)

def parse_md(md_text):
    """Extract structured data from the markdown research file."""
    data = {}

    # Extract company name from title
    m = re.search(r'^# (.+?) —', md_text, re.MULTILINE)
    data['company'] = m.group(1).strip() if m else 'COMPANY'

    # Extract date
    m = re.search(r'Generated: (.+)', md_text)
    data['date'] = m.group(1).strip() if m else ''

    # Extract copy-paste blocks (text inside ``` blocks after "Body (copy-paste):" or similar)
    def get_block(section_label):
        # Look for the label followed by a code block
        pattern = rf'{re.escape(section_label)}.*?```\n(.*?)```'
        m = re.search(pattern, md_text, re.DOTALL)
        return m.group(1).strip() if m else ''

    # Slide 1 boxes
    sections = md_text.split('**►')
    data['blocks'] = {}
    for sec in sections:
        if not sec.strip():
            continue
        # Get the label
        label_match = re.match(r'\s*(.+?)\*\*', sec)
        if not label_match:
            continue
        label = label_match.group(1).strip()
        # Get code blocks within this section
        code_blocks = re.findall(r'```\n(.*?)```', sec, re.DOTALL)
        # Get source lines
        sources = re.findall(r'Sources?:(.+?)(?:\n\n|\n\*\*|$)', sec, re.DOTALL)

        data['blocks'][label] = {
            'code_blocks': [b.strip() for b in code_blocks],
            'sources': [s.strip() for s in sources],
        }

    # Extract full sections for research backing
    backing_match = re.search(r'RESEARCH BACKING.*?─{10,}(.*?)(?:## SOURCES|$)', md_text, re.DOTALL)
    data['research_backing'] = backing_match.group(1).strip() if backing_match else ''

    # Extract sources
    sources_match = re.search(r'## SOURCES\n(.*?)$', md_text, re.DOTALL)
    data['sources_list'] = sources_match.group(1).strip() if sources_match else ''

    return data

def build_docx_from_md(md_path, docx_path):
    with open(md_path, 'r') as f:
        md_text = f.read()

    data = parse_md(md_text)
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # Title
    p = doc.add_heading(data['company'], level=0)
    for run in p.runs:
        run.font.size = Pt(24)
        run.font.color.rgb = DARK

    p = doc.add_paragraph()
    r = p.add_run(f"Stripe Sessions Slide Content — Copy-Paste Ready  |  {data['date']}")
    r.font.size, r.font.italic, r.font.color.rgb = Pt(10), True, GRAY

    # Process blocks in order
    blocks = data['blocks']

    # SLIDE 1
    add_slide_header(doc, 'SLIDE 1: "How ' + data['company'] + ' moves money today."')

    slide1_boxes = [
        ('BOX 01', 'Fragmented stacks'),
        ('BOX 02', 'Single point of failure'),
        ('BOX 03', 'Localization debt'),
        ('BOX 04', 'Reconciliation drag'),
    ]

    # Slide title
    for key, block_data in blocks.items():
        if 'SLIDE TITLE' in key and 'SLIDE 1' not in key and 'SLIDE 2' not in key:
            if 'moves money' in str(block_data.get('code_blocks', '')):
                add_copy_block(doc, "Slide Title", block_data['code_blocks'][0] if block_data['code_blocks'] else '')
                break

    add_separator(doc)

    for box_id, box_title in slide1_boxes:
        add_box_header(doc, box_id, box_title)
        # Find matching block
        for key, block_data in blocks.items():
            if box_id in key and box_title.lower() in key.lower():
                cbs = block_data['code_blocks']
                if len(cbs) >= 1:
                    add_copy_block(doc, "Title", cbs[0])
                if len(cbs) >= 2:
                    add_copy_block(doc, "Body", cbs[1])
                if block_data['sources']:
                    add_source(doc, f"Sources: {block_data['sources'][0]}")
                break
        add_separator(doc)

    # PSP Topology
    for key, block_data in blocks.items():
        if 'PSP TOPOLOGY' in key:
            add_box_header(doc, "PSP TOPOLOGY", "Data for the visual")
            for cb in block_data['code_blocks']:
                add_copy_block(doc, "PSPs", cb)
            if block_data['sources']:
                add_source(doc, f"Sources: {block_data['sources'][0]}")
            break

    add_separator(doc)

    # Blind Spots
    for key, block_data in blocks.items():
        if 'BLIND SPOTS' in key:
            add_box_header(doc, "BLIND SPOTS", "Labels for bottom row")
            for cb in block_data['code_blocks']:
                add_copy_block(doc, "APMs", cb)
            break

    doc.add_page_break()

    # SLIDE 2
    add_slide_header(doc, 'SLIDE 2: "From a processor mesh to one orchestration layer"')

    slide2_boxes = [
        ('BOX 01', 'Smart Routing'),
        ('BOX 02', 'Failover & Retries'),
        ('BOX 03', 'Local Payment Methods'),
        ('BOX 04', 'Unified Orchestration'),
    ]

    for box_id, box_title in slide2_boxes:
        add_box_header(doc, box_id, box_title)
        for key, block_data in blocks.items():
            if box_id in key and box_title.split()[0].lower() in key.lower():
                # Skip if this is a Slide 1 block
                if 'Fragmented' in key or 'failure' in key or 'Localization' in key or 'Reconciliation' in key:
                    continue
                cbs = block_data['code_blocks']
                if len(cbs) >= 1:
                    add_copy_block(doc, "Title", cbs[0])
                if len(cbs) >= 2:
                    add_copy_block(doc, "Body", cbs[1])
                break
        add_separator(doc)

    # Expected Impact
    for key, block_data in blocks.items():
        if 'EXPECTED IMPACT' in key:
            add_box_header(doc, "EXPECTED IMPACT", "Metric boxes")
            for cb in block_data['code_blocks']:
                add_copy_block(doc, "Metrics", cb)
            break

    add_separator(doc)

    # Success Cases
    for key, block_data in blocks.items():
        if 'SUCCESS CASES' in key:
            add_box_header(doc, "SUCCESS CASES", "Most relevant")
            for cb in block_data['code_blocks']:
                add_copy_block(doc, "Cases", cb)
            break

    doc.add_page_break()

    # Research Backing
    add_slide_header(doc, 'RESEARCH BACKING — Key Facts')

    if data['research_backing']:
        for line in data['research_backing'].split('\n'):
            line = line.strip()
            if not line:
                continue
            if line.startswith('**') and line.endswith('**'):
                p = doc.add_paragraph()
                r = p.add_run(line.replace('**', ''))
                r.font.bold, r.font.size = True, Pt(12)
            elif line.startswith('- ') or line.startswith('1. '):
                text = re.sub(r'\*\*(.+?)\*\*', r'\1', line.lstrip('- 0123456789. '))
                doc.add_paragraph(text, style='List Bullet')
            else:
                text = re.sub(r'\*\*(.+?)\*\*', r'\1', line)
                doc.add_paragraph(text)

    # Sources
    if data['sources_list']:
        doc.add_paragraph()
        p = doc.add_paragraph()
        r = p.add_run("SOURCES")
        r.font.bold, r.font.size = True, Pt(12)
        for line in data['sources_list'].split('\n'):
            line = line.strip()
            if line.startswith('- '):
                # Extract link text and URL
                link_match = re.match(r'- \[(.+?)\]\((.+?)\)', line)
                if link_match:
                    doc.add_paragraph(f"{link_match.group(1)}: {link_match.group(2)}", style='List Bullet')
                else:
                    doc.add_paragraph(line.lstrip('- '), style='List Bullet')

    doc.save(docx_path)
    print(f"Word document saved: {docx_path}")

if __name__ == '__main__':
    if len(sys.argv) < 2:
        print("Usage: python3 generate_stripe_sessions_docx.py <company-slug>")
        sys.exit(1)

    slug = sys.argv[1]
    # Check both flat and subfolder structures
    md_path = os.path.join(RESEARCH_DIR, f"{slug}-stripe-sessions-research.md")
    docx_path = os.path.join(RESEARCH_DIR, f"{slug}-stripe-sessions-research.docx")

    # Also check inside company subfolder (e.g., Research/Discord/)
    subfolder = os.path.join(RESEARCH_DIR, slug.capitalize())
    if not os.path.exists(md_path) and os.path.exists(subfolder):
        alt_md = os.path.join(subfolder, f"{slug}-stripe-sessions-research.md")
        if os.path.exists(alt_md):
            md_path = alt_md
            docx_path = os.path.join(subfolder, f"{slug}-stripe-sessions-research.docx")

    if not os.path.exists(md_path):
        print(f"Error: {md_path} not found")
        sys.exit(1)

    build_docx_from_md(md_path, docx_path)
