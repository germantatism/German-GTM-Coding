"""Build Salesforce meeting prep .docx for Yuno SDR."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT_PATH = "/Users/germantatis/Documents/GERMAN CLAUDE CODE/GTMCoding/Salesforce - Meeting Prep.docx"

YUNO_PURPLE = RGBColor(0x6B, 0x2C, 0xB0)
YUNO_DARK = RGBColor(0x1A, 0x1A, 0x2E)
YUNO_GRAY = RGBColor(0x5A, 0x5A, 0x6E)
YUNO_GREEN = RGBColor(0x10, 0x9A, 0x5C)
YUNO_RED = RGBColor(0xC4, 0x2B, 0x1C)
YUNO_AMBER = RGBColor(0xD4, 0x8A, 0x0E)

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

# Default font
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)
style.font.color.rgb = YUNO_DARK


def add_heading(text, level=1, color=YUNO_PURPLE, size=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.bold = True
    run.font.color.rgb = color
    if size is None:
        size = {1: 20, 2: 15, 3: 12}.get(level, 11)
    run.font.size = Pt(size)
    return p


def add_paragraph(text, bold=False, italic=False, color=None, size=11, space_after=6, indent=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.25
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return p


def add_bullet(text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(0.6 + level * 0.6)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.2
    run = p.add_run(text)
    run.font.size = Pt(11)
    return p


def add_mixed_bullet(parts, level=0):
    """parts: list of (text, bold) tuples."""
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(0.6 + level * 0.6)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.2
    for text, bold in parts:
        run = p.add_run(text)
        run.bold = bold
        run.font.size = Pt(11)
    return p


def shade_cell(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def add_table(headers, rows, header_bg="6B2CB0", widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.autofit = False
    table.alignment = WD_ALIGN_PARAGRAPH.LEFT

    if widths:
        for col_idx, w in enumerate(widths):
            for row in table.rows:
                row.cells[col_idx].width = Cm(w)

    # Header
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = ""
        shade_cell(hdr_cells[i], header_bg)
        p = hdr_cells[i].paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(10)

    # Rows
    for r_idx, row in enumerate(rows):
        cells = table.rows[r_idx + 1].cells
        for c_idx, val in enumerate(row):
            cells[c_idx].text = ""
            p = cells[c_idx].paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(10)
            if r_idx % 2 == 1:
                shade_cell(cells[c_idx], "F4EFFB")
    return table


def add_divider():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run("_" * 70)
    run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xD8)
    run.font.size = Pt(8)


def add_callout(label, body, color=YUNO_PURPLE):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Cm(0.4)
    run_label = p.add_run(f"{label}  ")
    run_label.bold = True
    run_label.font.color.rgb = color
    run_label.font.size = Pt(11)
    run_body = p.add_run(body)
    run_body.font.size = Pt(11)


# ============================================================
# COVER
# ============================================================
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.LEFT
title.paragraph_format.space_after = Pt(2)
r = title.add_run("yuno  |  Salesforce")
r.bold = True
r.font.size = Pt(11)
r.font.color.rgb = YUNO_PURPLE

cover_title = doc.add_paragraph()
cover_title.paragraph_format.space_before = Pt(30)
cover_title.paragraph_format.space_after = Pt(6)
r = cover_title.add_run("Preparación de Reunión")
r.bold = True
r.font.size = Pt(28)
r.font.color.rgb = YUNO_DARK

cover_sub = doc.add_paragraph()
cover_sub.paragraph_format.space_after = Pt(4)
r = cover_sub.add_run("Equipo Salesforce")
r.font.size = Pt(18)
r.font.color.rgb = YUNO_PURPLE

cover_date = doc.add_paragraph()
cover_date.paragraph_format.space_after = Pt(20)
r = cover_date.add_run("Fecha del brief: 2026-05-26  |  Preparado por: SDR Yuno")
r.font.size = Pt(10)
r.font.color.rgb = YUNO_GRAY
r.italic = True

# Briefing summary box
add_callout(
    "Tesis de la reunión:",
    "Salesforce opera un ecosistema multi-PSP fragmentado (Stripe + Adyen + Cybersource + cartridges) "
    "sin orquestación centralizada. El ángulo no es reemplazar PSPs, sino posicionar a Yuno como la capa "
    "de orquestación que conecta su estrategia de Agentic Commerce, su nueva integración nativa con Adyen "
    "y la expansión LATAM donde tienen entidades pero brechas operativas (México sin entidad local).",
    color=YUNO_PURPLE,
)

add_callout(
    "Disclaimer:",
    "Este documento se basa en research público con fuentes citadas. Algunos datos están marcados "
    'como "según datos web" o "asunciones realizadas" cuando la fuente primaria no pudo verificarse '
    "(según política de Yuno).",
    color=YUNO_AMBER,
)

doc.add_page_break()

# ============================================================
# TABLA DE CONTENIDOS
# ============================================================
add_heading("Tabla de Contenidos", level=1, size=18)
toc_items = [
    "1. Resumen Ejecutivo",
    "2. Quién es Salesforce — Contexto Crítico",
    "3. Stack de Pagos Actual",
    "4. Señales de Compra y Triggers",
    "5. Dolores Identificados (con evidencia)",
    "6. ÁNGULO DE VENTA — Qué les vendemos y cómo",
    "7. Mensajes Clave para la Reunión",
    "8. Preguntas de Discovery",
    "9. Objeciones Anticipadas",
    "10. Casos de Éxito Relevantes",
    "11. Próximos Pasos",
    "12. Apéndice — Fuentes Citadas",
]
for item in toc_items:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(item)
    run.font.size = Pt(11)

doc.add_page_break()

# ============================================================
# 1. RESUMEN EJECUTIVO
# ============================================================
add_heading("1. Resumen Ejecutivo", level=1, size=18)

add_paragraph(
    "Salesforce (FY26 revenue $41.5B, +10% YoY) es un B2B SaaS con presencia en 30+ países y un "
    "ecosistema de pagos fragmentado por producto. Su propia facturación corre principalmente vía "
    "wire/ACH a tres entidades regionales (SFDC Ireland para EMEA, Salesforce Inc. para AMER, entidad APAC). "
    "El flujo de tarjeta self-serve (<$5K USD) usa Salesforce Checkout con Stripe como backend (según datos web).",
    space_after=10,
)

add_paragraph(
    "La oportunidad para Yuno NO está en su facturación SaaS interna. Está en tres frentes:",
    bold=True,
    space_after=4,
)

add_mixed_bullet([
    ("Commerce Cloud merchants: ", True),
    ("su base de clientes B2C usa Stripe + Adyen + Cybersource + cartridges sin orquestación unificada.", False),
])
add_mixed_bullet([
    ("Agentic Commerce: ", True),
    ("alianza con Stripe + OpenAI (Oct 2025) y consejo UCP de Google (NRF 2026) expanden las superficies "
     "de pago, multiplicando la fragmentación.", False),
])
add_mixed_bullet([
    ("Expansión LATAM: ", True),
    ("tienen entidades en Argentina, Chile, Colombia y Brasil pero NO en México. Cross-border BIN flags "
     "es el primer escape de margen.", False),
])

add_paragraph("Score de prioridad: 17/22 puntos → 🔴 ALTA prioridad.", bold=True, color=YUNO_RED, space_after=14)

# Snapshot box
add_heading("Snapshot rápido", level=2, size=13)
snapshot_data = [
    ["Métrica", "Valor", "Fuente"],
]
snapshot_rows = [
    ["Revenue FY26", "$41.5B (+10% YoY)", "Salesforce Q4 FY26 earnings"],
    ["Clientes", "~150,000+", "Reportes públicos"],
    ["Países con entidad legal", "30+", "SEC FY25 Exhibit 21.1"],
    ["Top mercado #1", "USA (38.80% tráfico)", "SimilarWeb"],
    ["Top mercado #2", "Japón (23.74%)", "SimilarWeb"],
    ["Top mercado #3", "India (7.06%, +4.76% MoM)", "SimilarWeb"],
    ["PSPs identificados", "Stripe, Adyen, Cybersource + 5 más", "AppExchange + press releases"],
    ["Orquestador actual", "Ninguno (no evidencia pública)", "Search exhaustivo"],
    ["PCI DSS", "Level 1 Service Provider", "compliance.salesforce.com"],
    ["CFO actual", "Robin Washington (desde mar 2025)", "Press release SF feb 2025"],
    ["M&A reciente", "$8B Informatica (cerró nov 2025)", "Pulse2"],
    ["Tasa de falla cards (su ecosistema)", "~20% baseline en SF Billing", "Cyntexa (partner SF)"],
]
add_table(snapshot_data[0], snapshot_rows, widths=[5.5, 6.0, 5.0])

doc.add_page_break()

# ============================================================
# 2. QUIÉN ES SALESFORCE — CONTEXTO CRÍTICO
# ============================================================
add_heading("2. Quién es Salesforce — Contexto Crítico", level=1, size=18)

add_heading("2.1 Modelo de Negocio", level=2, size=13)
add_paragraph(
    "Salesforce es el líder mundial en CRM SaaS, pero NO es un merchant tradicional. Tienen dos "
    "perfiles de pago muy distintos que hay que separar mentalmente antes de la reunión:",
    space_after=6,
)
add_mixed_bullet([
    ("Salesforce como vendedor de su propio software: ", True),
    ("cobra a sus 150K clientes mediante wire/ACH (enterprise) y Salesforce Checkout (SMB self-serve, "
     "<$5K USD). Este NO es nuestro target principal.", False),
])
add_mixed_bullet([
    ("Salesforce como proveedor de Commerce Cloud: ", True),
    ("provee el stack de pagos a miles de retailers globales que SI son nuestro target. Aquí es donde "
     "Yuno como orquestador entra.", False),
])

add_heading("2.2 Estructura Legal Global", level=2, size=13)
add_paragraph(
    "Routing de facturación confirmado vía artículos de Salesforce Help (remittance):",
    space_after=6,
)
entity_data = [
    ["Región", "Entidad billing", "Monedas soportadas"],
]
entity_rows = [
    ["AMER", "Salesforce Inc.", "USD"],
    ["EMEA", "SFDC Ireland Limited", "USD, EUR, GBP"],
    ["EMEA (alt GBP)", "Salesforce UK Limited", "GBP, EUR"],
    ["APAC", "Entidad APAC dedicada", "Múltiples monedas locales"],
]
add_table(entity_data[0], entity_rows, widths=[3.5, 7.0, 6.0])

add_paragraph("", space_after=4)
add_callout(
    "GAP CRÍTICO LATAM:",
    "México NO aparece en el SEC FY25 Exhibit 21.1 de subsidiarias de Salesforce. Argentina, Chile, "
    "Colombia y Brasil sí. Esto significa que cobros mexicanos probablemente se enrutan vía la entidad "
    "US, generando cross-border BIN flags, mayores declines, y pérdida de APMs locales (OXXO, SPEI, MSI).",
    color=YUNO_RED,
)

add_heading("2.3 Mercados Top por Tráfico", level=2, size=13)
traffic_data = [
    ["Rank", "País", "% Tráfico", "Tendencia MoM", "Entidad local"],
]
traffic_rows = [
    ["1", "Estados Unidos", "38.80%", "+1.59%", "Sí"],
    ["2", "Japón", "23.74%", "+0.04%", "Sí"],
    ["3", "India", "7.06%", "+4.76%", "Sí"],
    ["4", "Reino Unido", "4.35%", "+5.58%", "Sí"],
    ["5", "Canadá", "2.51%", "+2.34%", "Sí"],
    ["6+", "Otros", "23.55%", "N/A", "Variable"],
]
add_table(traffic_data[0], traffic_rows, widths=[1.5, 4.5, 2.5, 3.0, 3.0])

add_paragraph(
    "Fuente: SimilarWeb (snapshot abril 2026). Datos detallados más allá del top 5 están detrás "
    'de paywall (según datos web).',
    italic=True,
    color=YUNO_GRAY,
    size=10,
    space_after=10,
)

add_heading("2.4 Producto: Lo que Salesforce vende relacionado a pagos", level=2, size=13)
add_mixed_bullet([
    ("Salesforce Billing (ahora Agentforce Revenue Management): ", True),
    ("herramienta de facturación y cobros para sus clientes. Conectores nativos: Cybersource, "
     "Authorize.Net, Payeezy, Payflow Pro.", False),
])
add_mixed_bullet([
    ("Commerce Cloud Payments: ", True),
    ("PSPs disponibles vía AppExchange: Stripe, Adyen, Braintree, Cybersource, PayTrace, Worldpay, "
     "Checkout.com, Authorize.Net.", False),
])
add_mixed_bullet([
    ("Salesforce Payments (Stripe-powered): ", True),
    ("solución embedida para merchants B2C, lanzada en 2021-2022.", False),
])
add_mixed_bullet([
    ("Salesforce Checkout: ", True),
    ("flujo self-serve para invoices <$5K USD del propio Salesforce.", False),
])

doc.add_page_break()

# ============================================================
# 3. STACK DE PAGOS ACTUAL
# ============================================================
add_heading("3. Stack de Pagos Actual", level=1, size=18)

add_heading("3.1 PSPs Confirmados en el Ecosistema", level=2, size=13)
psp_data = [
    ["PSP / Acquirer", "Rol en Salesforce", "Evidencia"],
]
psp_rows = [
    ["Wire / ACH (bancos regionales)", "Cobro principal B2B enterprise", "SF Help remittance pages"],
    ["Stripe", "Salesforce Checkout (SMB) + alianza ACP Oct 2025", "Stripe customer page + press release"],
    ["Cybersource (Visa)", "Default gateway en Salesforce Billing", "SF Help Billing Gateways"],
    ["Adyen", "Integración nativa GA en B2C Commerce 26.1 (2026)", "SF Help 26.1 release notes"],
    ["Braintree, PayTrace, Worldpay", "Cartridges en AppExchange para Commerce Cloud", "AppExchange listings"],
    ["PagBrasil", "Partner certificado para Pix + Boleto en Brasil", "PagBrasil partnership page"],
]
add_table(psp_data[0], psp_rows, widths=[5.0, 5.5, 6.0])

add_paragraph("", space_after=6)

add_heading("3.2 Estado de Orquestación", level=2, size=13)
add_callout(
    "Hallazgo clave:",
    "NO se encontró evidencia pública de que Salesforce use un orquestador (Spreedly, Primer, Gr4vy, "
    "CellPoint, APEXX, Yuno) para su propio billing o para su offering de Commerce Cloud. "
    "Gr4vy aparece en AppExchange como ofrecimiento a sus clientes, no como infra propia.",
    color=YUNO_GREEN,
)
add_paragraph(
    "Esto es nuestro principal punto de entrada. La fragmentación es un hecho documentado, no una "
    "asunción.",
    italic=True,
    space_after=10,
)

add_heading("3.3 PCI DSS y Manejo de Datos", level=2, size=13)
add_bullet("Nivel: PCI DSS Level 1 Service Provider (certificado desde Nov 2011, auditoría anual QSA)")
add_bullet("Cobertura: Salesforce Services + Hyperforce, Heroku, Agentforce/Einstein, B2C Commerce, "
           "Data Cloud, Marketing Cloud, Tableau Cloud")
add_bullet("Cardholder data: debe almacenarse SOLO en Encrypted Custom Fields (ECF); modelo de "
           "responsabilidad compartida (la certificación no se extiende automáticamente a instancias de clientes)")
add_bullet("Implicación: Yuno encaja bien con su modelo de tokenización; reduce scope PCI sin requerir "
           "que Salesforce expanda su propia certificación.")

doc.add_page_break()

# ============================================================
# 4. SEÑALES DE COMPRA Y TRIGGERS
# ============================================================
add_heading("4. Señales de Compra y Triggers", level=1, size=18)

add_paragraph(
    "Estos son los disparadores que justifican esta reunión AHORA y no en otro momento:",
    bold=True,
    space_after=8,
)

trigger_data = [
    ["#", "Trigger", "Fecha", "Por qué importa para Yuno"],
]
trigger_rows = [
    ["1",
     "Nuevo COFO Robin Washington (combina COO + CFO)",
     "Feb 2025 (efectivo mar 2025)",
     "Nuevos CFOs revisan stack de pagos en meses 6-18. A mayo 2026 está en mes 14. Ventana abierta."],
    ["2",
     "Cierre adquisición Informatica $8B",
     "Nov 18, 2025",
     "Integraciones post-M&A típicamente exponen fragmentación de infra. Momento de consolidación."],
    ["3",
     "Alianza Agentic Commerce Protocol con Stripe + OpenAI",
     "Oct 14, 2025",
     "Multiplica las superficies de pago. Orquestación pasa de nice-to-have a table stakes."],
    ["4",
     "Consejo Tech UCP de Google (con Amazon, Meta, Microsoft, Stripe)",
     "NRF 2026",
     "Salesforce ahora está en múltiples protocolos de pago agéntico simultáneamente."],
    ["5",
     "Adyen nativo GA en B2C Commerce 26.1",
     "2026",
     'Adyen pasa a ser "first class" junto a Stripe. Confirma que la estrategia es multi-PSP, no mono-PSP.'],
    ["6",
     "Expansión Hyperforce a 38+ regiones",
     "2023-2025",
     "Infraestructura global expandiéndose; cada nuevo mercado requiere local payments + compliance."],
    ["7",
     "Cyntexa (partner SF) publica ~20% baseline failure rate en SF Billing",
     "2025",
     "El propio ecosistema admite el problema. Es nuestro mejor opener."],
]
add_table(trigger_data[0], trigger_rows, widths=[0.8, 4.5, 3.2, 7.5])

doc.add_page_break()

# ============================================================
# 5. DOLORES IDENTIFICADOS
# ============================================================
add_heading("5. Dolores Identificados (con evidencia)", level=1, size=18)

add_paragraph(
    "Estos dolores están documentados en fuentes públicas. No son asunciones. Cada uno tiene un "
    "antídoto Yuno concreto.",
    space_after=10,
)

pain_data = [
    ["Dolor", "Evidencia", "Antídoto Yuno"],
]
pain_rows = [
    ["~20% de transacciones de tarjeta fallan en flujos de SF Billing",
     "Blog Cyntexa (partner SF), 2025",
     "Smart Routing +7% approval uplift, 50% recovery de declines"],
    ["Recuperación de pagos fallidos requiere intervención manual",
     "SF Help: blng_guidelines_resolving_failed_transactions",
     "Monitores en tiempo real (caso Rappi: ms vs 5-10 min manual)"],
    ["Account-locked-out cuando expira la tarjeta (cascada a collections)",
     "Trustpilot multiple cases 2025",
     "Card updater + dunning inteligente nativo en Yuno"],
    ["Klarna NO disponible para subscripciones / recurring en SF flows",
     "Documentación oficial de Klarna",
     "Yuno orquesta 1,000+ APMs incluyendo BNPL alternativos para recurring"],
    ["LATAM local methods requieren integraciones third-party (PagBrasil, etc.)",
     "PagBrasil partnership page",
     "Yuno trae Pix, Boleto, OXXO, SPEI, MSI nativos en una API"],
    ["México sin entidad local → cross-border BIN flags",
     "SEC FY25 Exhibit 21.1 (México no aparece)",
     "Yuno habilita PSPs locales en semanas (caso InDrive)"],
    ["Auto-renewal sin notificación + bloqueo de cancelación",
     "BBB + Trustpilot recurring complaints 2025-2026",
     'Mejora de UX en flujos de pago reduce churn por payment friction'],
    ["10% renewal price hike booked antes de hablar con cliente",
     "Salesflare blog 2025",
     "Transparencia de PSP economics permite renegociar margen"],
]
add_table(pain_data[0], pain_rows, widths=[6.0, 5.0, 6.0])

doc.add_page_break()

# ============================================================
# 6. ÁNGULO DE VENTA
# ============================================================
add_heading("6. ÁNGULO DE VENTA — Qué les vendemos y cómo", level=1, size=18, color=YUNO_PURPLE)

add_callout(
    "Tesis Central:",
    'Yuno NO compite con Stripe ni con Adyen ni con Cybersource. Yuno es la capa de orquestación '
    'que hace que esos PSPs coexistan operativamente. Es el "cerebro de routing" sobre la fragmentación '
    "que Salesforce ya tiene.",
    color=YUNO_PURPLE,
)

add_heading("6.1 Posicionamiento (Cómo nos presentamos)", level=2, size=13)
add_paragraph(
    'Frase de apertura recomendada:',
    bold=True,
    space_after=4,
)
add_paragraph(
    '"Yuno es la capa de orquestación que conecta Stripe, Adyen, Cybersource y cualquier otro PSP en '
    'una sola API. No reemplazamos a sus PSPs actuales. Hacemos que sumen, en lugar de sumar complejidad."',
    italic=True,
    color=YUNO_PURPLE,
    space_after=10,
    indent=0.6,
)

add_heading("6.2 Los 3 Frentes de Venta", level=2, size=13)

# Frente 1
add_paragraph("FRENTE 1: Commerce Cloud Merchants (el más grande)", bold=True, color=YUNO_PURPLE, space_after=4)
add_mixed_bullet([
    ("Qué les vendemos: ", True),
    ("Yuno como AppExchange offering para sus merchants B2C. Una integración, todos los PSPs.", False),
])
add_mixed_bullet([
    ("Por qué les conviene: ", True),
    ("Sus merchants ya pelean con cartridges múltiples. Una integración Yuno = un solo punto de "
     "mantenimiento y soporte para Salesforce. Mejor experiencia para sus clientes.", False),
])
add_mixed_bullet([
    ("Cómo monetizamos: ", True),
    ("Revenue share por merchant activado, o referencia con incentivo. Modelo de partner clásico de AppExchange.", False),
])
add_mixed_bullet([
    ("Caso de éxito a citar: ", True),
    ("Rappi (zero implementation time, 80% menos resolución de analistas).", False),
])

add_paragraph("", space_after=6)

# Frente 2
add_paragraph("FRENTE 2: Agentforce Commerce / Agentic Payments", bold=True, color=YUNO_PURPLE, space_after=4)
add_mixed_bullet([
    ("Qué les vendemos: ", True),
    ("Yuno como vault unificado y routing engine para todos los flujos agénticos (ACP con Stripe, "
     "UCP con Google).", False),
])
add_mixed_bullet([
    ("Por qué les conviene: ", True),
    ("Cada protocolo nuevo (ACP, UCP, próximos) multiplica los puntos donde tokenizar y enrutar. "
     "Yuno abstrae esa complejidad detrás de una API estable.", False),
])
add_mixed_bullet([
    ("Pitch específico: ", True),
    ('"En agentic commerce, el customer journey no termina en checkout. Empieza en un asistente. '
     'Necesitan orquestación que viva más arriba que el PSP."', False),
])
add_mixed_bullet([
    ("Caso de éxito a citar: ", True),
    ("Reserva (+4% approval en <3 meses con multi-PSP routing).", False),
])

add_paragraph("", space_after=6)

# Frente 3
add_paragraph("FRENTE 3: Expansión LATAM (gap México + optimización resto)", bold=True, color=YUNO_PURPLE, space_after=4)
add_mixed_bullet([
    ("Qué les vendemos: ", True),
    ("Yuno como motor de pagos LATAM-nativo. MoR coverage + local APMs + smart routing local.", False),
])
add_mixed_bullet([
    ("Por qué les conviene: ", True),
    ("Tienen entidades en AR, CL, CO, BR pero México es un gap. Cards mexicanas vía US merchant ID "
     "se queman en declines. Y en los países donde sí tienen entidad, no necesariamente tienen el "
     "PSP local óptimo.", False),
])
add_mixed_bullet([
    ("Pitch específico: ", True),
    ('"InDrive entró a 10 mercados LATAM en menos de 8 meses con Yuno, manteniendo 90% approval. '
     'Mismo modelo aplicable a su huella LATAM."', False),
])
add_mixed_bullet([
    ("Caso de éxito a citar: ", True),
    ("InDrive (10 LATAM <8 meses, 90% approval, 4.5% recovery)."
     " Livelo (+5% approval, 50% recovery).", False),
])

add_paragraph("", space_after=10)

add_heading("6.3 Por Qué Salesforce Compraría Yuno (lógica fría)", level=2, size=13)
add_mixed_bullet([
    ("Robin Washington (COFO) ", True),
    ("está en el ciclo de revisión de procurement de finanzas / pagos. Yuno entra a tiempo.", False),
])
add_mixed_bullet([
    ("Post-Informatica ", True),
    ("hay momentum de consolidación. Yuno consolida la capa de routing de pagos.", False),
])
add_mixed_bullet([
    ("Adyen nativo en 26.1 ", True),
    ("confirma estrategia multi-PSP. Yuno orquesta multi-PSP. La narrativa se alinea, no se contradice.", False),
])
add_mixed_bullet([
    ("PCI DSS Level 1 ", True),
    ("ya cubierto en Yuno. No expande su scope de compliance.", False),
])
add_mixed_bullet([
    ("Cero churn ", True),
    ("de PSPs actuales. Yuno es aditivo, no sustitutivo. Reduce fricción política interna.", False),
])

add_heading("6.4 Cuál NO es el ángulo (errores a evitar)", level=2, size=13)
add_mixed_bullet([
    ("NO ", True),
    ('vender "reemplaza tu PSP". Salesforce tiene partnerships estratégicas con Stripe y Adyen.', False),
], )
add_mixed_bullet([
    ("NO ", True),
    ("decir que Salesforce Billing no funciona. Es su producto, es su orgullo. Hablar de routing inteligente "
     "encima de su producto, no en contra.", False),
])
add_mixed_bullet([
    ("NO ", True),
    ("hablar mal de Cybersource. Es de Visa. Visa es socio de muchas cosas.", False),
])
add_mixed_bullet([
    ("NO ", True),
    ("usar lenguaje de displacement. Usar lenguaje de orchestration / optimization layer.", False),
])

doc.add_page_break()

# ============================================================
# 7. MENSAJES CLAVE PARA LA REUNIÓN
# ============================================================
add_heading("7. Mensajes Clave para la Reunión", level=1, size=18)

add_heading("7.1 Apertura (primeros 60 segundos)", level=2, size=13)
add_paragraph(
    '"Gracias por el tiempo. Antes de entrar a Yuno, quería compartirles que estuvimos analizando su '
    'huella de pagos y vimos 3 cosas interesantes: la alianza ACP con Stripe + OpenAI de octubre, '
    'la integración nativa de Adyen en B2C Commerce 26.1, y que México no aparece en su roster de entidades '
    'del SEC. Cualquiera de esos tres temas merece una conversación. Por dónde quieren empezar?"',
    italic=True,
    color=YUNO_DARK,
    space_after=10,
    indent=0.6,
)

add_heading("7.2 Bullet Points para Defender", level=2, size=13)
add_bullet("Yuno = single API para 1,000+ payment methods en 200+ países")
add_bullet("Smart Routing → +7% approval uplift promedio")
add_bullet("Recovery de declines: 50% (Livelo case)")
add_bullet("Time to new market: semanas, no meses (InDrive: 10 LATAM en <8 meses)")
add_bullet("Real-time monitoring (Rappi: ms detection vs 5-10 min manual)")
add_bullet("No-code PSP enablement → el equipo de payments puede experimentar sin sprint de ingeniería")
add_bullet("PCI DSS Level 1 → no expande scope de compliance del cliente")
add_bullet("MoR coverage en LATAM → habilita cobros locales sin entidad legal local")

add_heading("7.3 Datos para Citar (todos verificados)", level=2, size=13)
add_mixed_bullet([
    ("Stat 1: ", True),
    ('"Su propio partner Cyntexa publicó que ~20% de las transacciones con tarjeta fallan en flujos '
     'de Salesforce Billing." Fuente: blog público de Cyntexa, 2025.', False),
])
add_mixed_bullet([
    ("Stat 2: ", True),
    ('"México no aparece en el SEC FY25 Exhibit 21.1 de subsidiarias de Salesforce. Argentina, Chile, '
     'Colombia y Brasil sí." Fuente: SEC EDGAR público.', False),
])
add_mixed_bullet([
    ("Stat 3: ", True),
    ('"En FY26 Salesforce hizo $41.5B y Agentforce está en $800M ARR creciendo +169% YoY. Las '
     'superficies de pago se están multiplicando con agentic." Fuente: Salesforce Q4 FY26 earnings.', False),
])
add_mixed_bullet([
    ("Stat 4: ", True),
    ('"Klarna no está disponible para subscripciones recurrentes en sus flujos, según la propia '
     'documentación de Klarna. Eso es una gap para SaaS recurring." Fuente: documentación oficial Klarna.', False),
])

doc.add_page_break()

# ============================================================
# 8. PREGUNTAS DE DISCOVERY
# ============================================================
add_heading("8. Preguntas de Discovery", level=1, size=18)

add_paragraph(
    "El objetivo de estas preguntas es validar dolor real, no asumir. Diseñadas para que ellos "
    "hablen, no nosotros.",
    space_after=10,
)

add_heading("8.1 Sobre su estrategia multi-PSP", level=2, size=13)
add_bullet('"Con la integración nativa de Adyen en 26.1, cómo están pensando la coexistencia con Stripe '
           'para sus merchants de Commerce Cloud?"')
add_bullet('"Los merchants pueden activar múltiples PSPs simultáneamente, o tienen que elegir uno?"')
add_bullet('"Tienen routing dinámico entre PSPs, o el merchant define el PSP por defecto y eso es?"')

add_heading("8.2 Sobre LATAM", level=2, size=13)
add_bullet('"Cómo manejan hoy los cobros de clientes mexicanos sin entidad local en México?"')
add_bullet('"PagBrasil cubre Brasil. Tienen partners equivalentes para Pix en otros corredores o '
           'para OXXO/SPEI en México?"')
add_bullet('"Cuál es el roadmap de expansión LATAM en FY27?"')

add_heading("8.3 Sobre Agentic Commerce", level=2, size=13)
add_bullet('"Cómo están abstrayendo la tokenización entre ACP (Stripe) y UCP (Google)? Es un vault '
           'unificado o uno por protocolo?"')
add_bullet('"Quién es el dueño interno de la estrategia de payments dentro de Agentforce?"')
add_bullet('"Ven los flujos agénticos como una extensión de Commerce Cloud o como una capa nueva?"')

add_heading("8.4 Sobre el dolor operativo", level=2, size=13)
add_bullet('"Cyntexa publicó que ~20% de las transacciones fallan en SF Billing. Esa cifra coincide '
           'con lo que ven internamente?"')
add_bullet('"Cómo recuperan las transacciones falladas hoy? Es proceso manual o automatizado?"')
add_bullet('"Quién en el equipo está midiendo approval rate por PSP, por mercado, por BIN?"')

add_heading("8.5 Sobre toma de decisión", level=2, size=13)
add_bullet('"Bajo el nuevo COFO Robin Washington, ha cambiado el approach a vendor consolidation '
           'en finanzas / payments?"')
add_bullet('"Una integración como Yuno entraría como AppExchange listing, como partnership directo, '
           'o como infra interna?"')
add_bullet('"Qué tendría que pasar para que sea una conversación de procurement seria en los próximos '
           '90 días?"')

doc.add_page_break()

# ============================================================
# 9. OBJECIONES ANTICIPADAS
# ============================================================
add_heading("9. Objeciones Anticipadas y Cómo Responderlas", level=1, size=18)

objections = [
    {
        "obj": '"Ya tenemos relación estratégica con Stripe y Adyen. No necesitamos otro PSP."',
        "resp": "Correcto, y eso es precisamente por qué Yuno tiene sentido. NO somos un PSP. Somos la "
                "capa de orquestación que hace que Stripe Y Adyen Y Cybersource coexistan en una sola API. "
                "Su relación con esos PSPs no cambia. Lo que cambia es la complejidad operativa de "
                "mantenerlos a todos en paralelo."
    },
    {
        "obj": '"Salesforce Billing y Revenue Cloud ya hacen routing. No necesitamos un tercero."',
        "resp": "Salesforce Billing maneja invoicing y subscription logic, lo cual es excelente. El gap "
                "es smart routing real-time entre múltiples PSPs y APMs, especialmente cuando entran "
                "flujos agénticos. Yuno no reemplaza Billing. Se conecta como la capa de execution "
                "donde Billing genera la intención de pago."
    },
    {
        "obj": '"No tenemos ancho de banda de ingeniería para integrar otra cosa."',
        "resp": "Por eso construimos no-code PSP enablement. Caso InDrive: 10 mercados LATAM en menos "
                "de 8 meses, con cero líneas de código nuevas por mercado adicional. Y el modelo de "
                "AppExchange listing hace que el go-to-market sea casi cero esfuerzo para Salesforce."
    },
    {
        "obj": '"México no es prioridad para nosotros."',
        "resp": "Entendido. México fue un ejemplo del patrón. El patrón es: cualquier mercado donde "
                "Salesforce expanda en los próximos 24 meses (FY27), Yuno habilita el lift en semanas "
                "vs trimestres. Qué mercados sí están en roadmap?"
    },
    {
        "obj": '"Yuno es muy nuevo / no los conocemos."',
        "resp": "Operamos con InDrive, Rappi, Livelo, Reserva, ya processando volumen significativo "
                "en LATAM y expandiéndonos globalmente. Podemos hacer una intro con cualquiera de "
                "ellos como referencia. Y trabajamos con PCI DSS Level 1, igual que ustedes."
    },
    {
        "obj": '"Esto suena más para nuestros merchants que para Salesforce mismo."',
        "resp": "Exacto, y ese es el frente principal. Estamos proponiendo dos modelos: (1) AppExchange "
                "listing donde sus merchants activan Yuno como capa de orquestación, y (2) revenue "
                "share entre Salesforce y Yuno por cada merchant activado. Win-win."
    },
    {
        "obj": '"Cuál es el costo?"',
        "resp": "Yuno es modelo SaaS + bps por transacción. La conversación de pricing tiene sentido "
                "cuando definimos scope (qué frentes activar). Antes de eso, la pregunta de valor es: "
                "+7% approval uplift sobre el GMV procesado, cuánto vale eso para sus merchants?"
    },
]

for i, o in enumerate(objections, 1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(f"Objeción {i}: ")
    r.bold = True
    r.font.color.rgb = YUNO_RED
    r.font.size = Pt(11)
    r2 = p.add_run(o["obj"])
    r2.italic = True
    r2.font.size = Pt(11)

    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(4)
    p2.paragraph_format.left_indent = Cm(0.6)
    r = p2.add_run("Respuesta: ")
    r.bold = True
    r.font.color.rgb = YUNO_GREEN
    r.font.size = Pt(11)
    r2 = p2.add_run(o["resp"])
    r2.font.size = Pt(11)

doc.add_page_break()

# ============================================================
# 10. CASOS DE ÉXITO RELEVANTES
# ============================================================
add_heading("10. Casos de Éxito Relevantes", level=1, size=18)

add_paragraph(
    "Cuatro casos para tener en boca durante la reunión. Cada uno ataca un dolor distinto.",
    space_after=10,
)

cases = [
    {
        "name": "InDrive",
        "color": YUNO_PURPLE,
        "stat": "10 mercados LATAM activados en <8 meses, 90% approval rate, 4.5% recovery",
        "when": "Citar cuando hablen de expansión geográfica o entrada a nuevos mercados",
        "pitch": "Si Salesforce planea acelerar LATAM en FY27, este es el modelo. Cero código nuevo "
                 "por mercado, time-to-live medido en semanas.",
    },
    {
        "name": "Rappi",
        "color": YUNO_PURPLE,
        "stat": "Zero implementation time, 80% menos resolución de tickets por analistas, "
                "detección de problemas en milisegundos vs 5-10 min manual",
        "when": "Citar cuando hablen del 20% failure rate o de overhead operativo",
        "pitch": "El equipo de payments de Salesforce probablemente está dedicando horas-persona a "
                 "investigar declines. Real-time monitoring elimina esa carga.",
    },
    {
        "name": "Livelo",
        "color": YUNO_PURPLE,
        "stat": "+5% approval rate, 50% recovery de transacciones falladas",
        "when": "Citar cuando hablen de optimization de funnel de checkout",
        "pitch": "5 puntos de approval sobre el GMV de un Commerce Cloud merchant grande es revenue "
                 "incremental sin más gasto de marketing.",
    },
    {
        "name": "Reserva",
        "color": YUNO_PURPLE,
        "stat": "+4% approval rate en menos de 3 meses",
        "when": "Citar cuando duden del time-to-value",
        "pitch": "Time-to-value de 12 semanas en producción, con resultados medibles antes del trimestre.",
    },
]

for c in cases:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(f"{c['name']}")
    r.bold = True
    r.font.size = Pt(14)
    r.font.color.rgb = c["color"]

    p2 = doc.add_paragraph()
    p2.paragraph_format.left_indent = Cm(0.4)
    p2.paragraph_format.space_after = Pt(3)
    r = p2.add_run("Stat clave: ")
    r.bold = True
    r.font.size = Pt(11)
    r2 = p2.add_run(c["stat"])
    r2.font.size = Pt(11)

    p3 = doc.add_paragraph()
    p3.paragraph_format.left_indent = Cm(0.4)
    p3.paragraph_format.space_after = Pt(3)
    r = p3.add_run("Cuándo usarlo: ")
    r.bold = True
    r.font.size = Pt(11)
    r2 = p3.add_run(c["when"])
    r2.font.size = Pt(11)

    p4 = doc.add_paragraph()
    p4.paragraph_format.left_indent = Cm(0.4)
    p4.paragraph_format.space_after = Pt(6)
    r = p4.add_run("Cómo enmarcarlo: ")
    r.bold = True
    r.font.size = Pt(11)
    r2 = p4.add_run(c["pitch"])
    r2.font.size = Pt(11)

doc.add_page_break()

# ============================================================
# 11. PRÓXIMOS PASOS
# ============================================================
add_heading("11. Próximos Pasos Sugeridos", level=1, size=18)

add_paragraph(
    "Dependiendo de cómo evolucione la conversación, estos son los CTAs en orden de preferencia:",
    space_after=8,
)

next_data = [
    ["#", "Acción", "Quién", "Timeline"],
]
next_rows = [
    ["1",
     "Workshop técnico de 60 min con el equipo de payments + arquitectura de Salesforce",
     "SDR Yuno + Solution Engineer + Salesforce payments lead",
     "Próximas 2 semanas"],
    ["2",
     "POC sandbox: integrar Yuno + Stripe + Adyen en un endpoint de prueba de Commerce Cloud",
     "Solution Engineer Yuno + Salesforce Commerce Cloud dev",
     "Próximas 4 semanas"],
    ["3",
     "Intro call con InDrive (referencia LATAM) o Rappi (referencia operacional)",
     "AE Yuno coordina",
     "Próximas 3 semanas"],
    ["4",
     "Conversación sobre AppExchange listing como partner integration",
     "BD Yuno + Salesforce ISV partnerships team",
     "Próximas 6 semanas"],
    ["5",
     "Si hay alineación: business case formal con números específicos sobre GMV de Commerce Cloud",
     "AE Yuno + Salesforce procurement",
     "60-90 días"],
]
add_table(next_data[0], next_rows, widths=[0.8, 6.5, 5.0, 3.5])

add_paragraph("", space_after=8)
add_callout(
    "Objetivo de la reunión:",
    "Salir con (a) un follow-up calendarizado, (b) nombres de stakeholders adicionales del lado de "
    "Salesforce (CTO de Commerce Cloud, líder de Agentforce payments, dueño de strategy en LATAM) "
    "y (c) confirmación del frente prioritario (Commerce Cloud merchants vs Agentic vs LATAM).",
    color=YUNO_PURPLE,
)

doc.add_page_break()

# ============================================================
# 12. APÉNDICE — FUENTES
# ============================================================
add_heading("12. Apéndice — Fuentes Citadas", level=1, size=18)

add_paragraph(
    "Todas las afirmaciones de este documento provienen de las siguientes fuentes públicas. "
    'Items marcados "según datos web" o "asunción" están señalados explícitamente.',
    italic=True,
    color=YUNO_GRAY,
    size=10,
    space_after=10,
)

sources = [
    ("Tráfico y mercados", [
        "SimilarWeb — salesforce.com: https://www.similarweb.com/website/salesforce.com/",
        "SimilarWeb — AppExchange subdomain: https://www.similarweb.com/website/appexchange.salesforce.com/",
    ]),
    ("Subsidiarias y entidades legales", [
        "SEC FY25 Exhibit 21.1: https://www.sec.gov/Archives/edgar/data/1108524/000110852425000006/ex211listofsubsidiariesfy25.htm",
        "LEI SFDC Ireland Ltd: https://www.leikart.com/leicert/549300UONRKY537S9629/",
    ]),
    ("Routing de facturación (Salesforce Help)", [
        "EMEA EUR: https://help.salesforce.com/s/articleView?id=000350632",
        "EMEA USD: https://help.salesforce.com/s/articleView?id=000350631",
        "EMEA GBP: https://help.salesforce.com/s/articleView?id=000350634",
        "AMER: https://help.salesforce.com/s/articleView?id=000349348",
        "APAC: https://help.salesforce.com/s/articleView?id=000340087",
    ]),
    ("PSPs y stack de pagos", [
        "Salesforce Checkout (campaign): https://www.salesforce.com/campaign/salesforce-checkout/",
        "SF Billing Payment Gateways: https://help.salesforce.com/s/articleView?id=sales.blng_payment_gateways_intro.htm",
        "Commerce Cloud Payments AppExchange: https://appexchange.salesforce.com/appxListingDetail?listingId=a0N4V00000H0jIVUAZ",
        "Stripe customer page (Salesforce): https://stripe.com/customers/salesforce",
        "Stripe + OpenAI + Salesforce ACP (Oct 14 2025): https://www.salesforce.com/news/press-releases/2025/10/14/stripe-openai-agentic-commerce-protocol-announcement/",
        "B2C Commerce 26.1 Adyen native: https://help.salesforce.com/s/articleView?id=commerce.b2c_rn_26_1_adyen.htm",
        "SF Blog Adyen native: https://www.salesforce.com/blog/salesforce-native-integration-adyen/",
        "Adyen cartridge GitHub: https://github.com/Adyen/adyen-salesforce-commerce-cloud/releases",
        "Klarna BNPL recurring restriction: https://documentation.b2c.commercecloud.salesforce.com/DOC3/topic/com.demandware.dochelp/content/b2c_commerce/topics/salesforce_payments/b2c_buy_now_pay_later_with_klarna.html",
        "PagBrasil partnership: https://www.pagbrasil.com/integration/app-salesforce/",
    ]),
    ("Complaints y operación", [
        "BBB Salesforce: https://www.bbb.org/us/ca/san-francisco/profile/computer-hardware/salesforcecom-inc-1116-44164/complaints",
        "Trustpilot Salesforce: https://www.trustpilot.com/review/salesforce.com",
        "Salesflare blog (MSA, renewal): https://blog.salesflare.com/salesforce-contracts",
        "Cyntexa ~20% baseline failure: https://products.cyntexa.com/blog/failed-payment-in-salesforce/",
        "SF Help failed transactions: https://help.salesforce.com/s/articleView?id=sf.blng_guidelines_resolving_failed_transactions.htm",
    ]),
    ("Liderazgo y M&A", [
        "Own Company acquisition (Sep 2024): https://www.salesforce.com/news/press-releases/2024/09/05/salesforce-signs-definitive-agreement-to-acquire-own-company/",
        "Robin Washington COFO appointment (Feb 2025): https://www.salesforce.com/news/press-releases/2025/02/05/salesforce-appoints-robin-washington-as-president-and-chief-operating-and-financial-officer/",
        "Informatica acquisition agreement (May 2025): https://www.salesforce.com/news/press-releases/2025/05/27/salesforce-signs-definitive-agreement-to-acquire-informatica/",
        "Informatica close (Nov 2025): https://pulse2.com/salesforce-completes-8-billion-acquisition-of-informatica-to-accelerate-agentic-ai-data-platform/",
        "Hyperforce expansion: https://www.salesforceben.com/5-reasons-to-pay-attention-to-salesforce-hyperforce-in-2025/",
    ]),
    ("Financieros", [
        "Salesforce Q4 FY26 earnings: https://www.salesforce.com/news/press-releases/2026/02/25/fy26-q4-earnings/",
        "Salesforce FY26 10-K ARS: https://www.sec.gov/Archives/edgar/data/1108524/000114036126015165/ef20070692_ars.pdf",
    ]),
    ("PCI DSS", [
        "Salesforce Compliance Portal: https://compliance.salesforce.com/en/categories/pci-dss",
        "AoC Salesforce Services: https://compliance.salesforce.com/en/documents/a005A00000k4aybQAA",
        "AoC Data Cloud/Agentforce: https://compliance.salesforce.com/en/documents/a006e00001ANWcGAAX",
    ]),
    ("Agentic Commerce", [
        "Stripe blog NRF 2026 trends: https://stripe.com/blog/three-agentic-commerce-trends-nrf-2026",
        "Agentforce Commerce announcement: https://www.salesforce.com/news/stories/agentforce-commerce-capabilities-announcement/",
    ]),
]

for category, urls in sources:
    add_heading(category, level=3, size=12)
    for url in urls:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.4)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run("• " + url)
        run.font.size = Pt(9)
        run.font.color.rgb = YUNO_GRAY

# Footer note
add_divider()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Documento preparado por SDR Yuno | Brief base: salesforce-2026-05-26.md | "
                "Última actualización: 2026-05-26")
run.italic = True
run.font.size = Pt(9)
run.font.color.rgb = YUNO_GRAY

doc.save(OUTPUT_PATH)
print(f"OK: {OUTPUT_PATH}")
